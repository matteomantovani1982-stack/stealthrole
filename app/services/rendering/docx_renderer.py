"""
app/services/rendering/docx_renderer.py

DOCX renderer: applies an EditPlan to the original CV DOCX.

Core constraint: The original DOCX is the source of truth for formatting.
Claude's EditPlan contains ONLY content instructions — never layout.
This module bridges the gap: it reads the EditPlan and applies each
operation to the document using python-docx, preserving all original
formatting (fonts, sizes, spacing, margins, styles, tables).

Operations supported (matching EditOperation enum):
  replace_text  — replace entire paragraph text, preserve run formatting
  replace_run   — replace one specific run's text, preserve its formatting
  insert_after  — insert new paragraph after a given index
  delete        — remove paragraph at index
  restyle       — change paragraph style name only

Headline/summary edits:
  Handled separately via HeadlineSummaryEdit — finds the first non-empty
  paragraphs and applies the new headline/summary content.

Design principles:
  1. Never change fonts, sizes, margins, or spacing — only text content
  2. Apply edits in reverse index order (high → low) so indices stay valid
  3. Copy run formatting from the original run when replacing text
  4. Preserve empty paragraphs (they define visual spacing)
  5. All operations are idempotent — safe to re-apply on retry

Limitations (future work):
  - Tables: content inside table cells is not edited (Sprint 2)
  - Text boxes: not supported by python-docx
  - Tracked changes: stripped on save (acceptable for CV use case)
"""

import copy
import io
import uuid
from dataclasses import dataclass
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

import structlog

logger = structlog.get_logger(__name__)


# ── Formatting copy helpers ───────────────────────────────────────────────────

def _copy_run_format(source_run: Run, target_run: Run) -> None:
    """
    Copy all formatting from source_run to target_run.

    Uses deep copy of the underlying XML element's rPr (run properties)
    to capture everything: font, size, bold, italic, colour, spacing, etc.
    This is the safest approach — copying individual attributes would
    miss obscure properties (kerning, character spacing, etc.)
    """
    # Get source run properties element
    source_rpr = source_run._r.find(qn("w:rPr"))
    if source_rpr is None:
        return  # Source has no explicit formatting — nothing to copy

    # Remove existing target run properties
    target_rpr = target_run._r.find(qn("w:rPr"))
    if target_rpr is not None:
        target_run._r.remove(target_rpr)

    # Deep copy and insert source properties into target
    new_rpr = copy.deepcopy(source_rpr)
    target_run._r.insert(0, new_rpr)


def _copy_paragraph_format(source_para: Paragraph, target_para: Paragraph) -> None:
    """
    Copy paragraph-level formatting (pPr) from source to target.
    Covers: alignment, spacing, indentation, list style, outline level.
    """
    source_ppr = source_para._p.find(qn("w:pPr"))
    if source_ppr is None:
        return

    target_ppr = target_para._p.find(qn("w:pPr"))
    if target_ppr is not None:
        target_para._p.remove(target_ppr)

    new_ppr = copy.deepcopy(source_ppr)
    target_para._p.insert(0, new_ppr)


def _get_first_non_empty_run(para: Paragraph) -> Run | None:
    """Return the first run with non-empty text, or None."""
    for run in para.runs:
        if run.text.strip():
            return run
    return None


# ── Core edit appliers ────────────────────────────────────────────────────────

def _apply_replace_text(para: Paragraph, new_text: str) -> None:
    """
    Replace all text in a paragraph while preserving the formatting
    of the first non-empty run.

    Strategy:
    1. Capture first run's formatting (the dominant formatting)
    2. Clear all existing runs
    3. Add a single new run with the new text and copied formatting

    Why single run? The LLM gives us one text string, not a list of runs.
    Splitting into multiple runs would require re-inferring formatting boundaries,
    which is error-prone. Single run is safe and visually correct for most cases.
    """
    if not para.runs:
        # No runs at all — just set via text property
        para.text = new_text
        return

    # Capture formatting from first non-empty run (or first run if all empty)
    reference_run = _get_first_non_empty_run(para) or para.runs[0]

    # Clear all existing runs by removing their XML elements
    for run in para.runs:
        run._r.getparent().remove(run._r)

    # Add single new run with new text
    new_run = para.add_run(new_text)
    _copy_run_format(reference_run, new_run)


def _apply_replace_run(para: Paragraph, run_index: int, new_text: str) -> None:
    """
    Replace text of a specific run by index, preserving its formatting.
    Silently skips if run_index is out of bounds.
    """
    runs = para.runs
    if run_index >= len(runs):
        logger.warning(
            "render_replace_run_index_out_of_bounds",
            extra={
                "run_index": run_index,
                "total_runs": len(runs),
            },
        )
        return

    target_run = runs[run_index]
    original_text = target_run.text
    target_run.text = new_text
    logger.debug(
        "render_run_replaced",
        extra={"from": original_text[:50], "to": new_text[:50]},
    )


def _apply_restyle(para: Paragraph, style_name: str, doc: Document) -> None:
    """
    Change a paragraph's style name.
    Falls back to 'Normal' if the style doesn't exist in the document.
    """
    available_styles = {s.name for s in doc.styles}
    if style_name in available_styles:
        para.style = doc.styles[style_name]
    else:
        logger.warning(
            "render_style_not_found",
            extra={"requested": style_name, "fallback": "Normal"},
        )
        para.style = doc.styles["Normal"]


# ── Paragraph index utilities ─────────────────────────────────────────────────

def _get_all_body_paragraphs(doc: Document) -> list[Paragraph]:
    """
    Return all body paragraphs in document order.

    Note: This returns only top-level body paragraphs.
    Paragraphs inside tables, text boxes, headers/footers
    are excluded — matching what the parser extracted.
    """
    return list(doc.paragraphs)


def _insert_paragraph_after(doc: Document, after_index: int, new_text: str, style_name: str | None = None) -> Paragraph:
    """
    Insert a new paragraph after the paragraph at after_index.

    Uses direct XML manipulation because python-docx's add_paragraph()
    always appends to the end. We need insertion at an arbitrary position.

    Args:
        doc:         The Document object
        after_index: Zero-based index of the paragraph to insert after
        new_text:    Text content for the new paragraph
        style_name:  Optional style name; defaults to 'Normal'

    Returns:
        The newly inserted Paragraph object
    """
    paragraphs = _get_all_body_paragraphs(doc)

    if after_index >= len(paragraphs):
        # If index is past the end, append to document
        return doc.add_paragraph(new_text, style=style_name or "Normal")

    # Get the reference paragraph's XML element
    ref_para = paragraphs[after_index]
    ref_element = ref_para._p

    # Create new paragraph XML element
    new_para_element = OxmlElement("w:p")

    # Add paragraph properties if style specified
    if style_name:
        ppr = OxmlElement("w:pPr")
        pstyle = OxmlElement("w:pStyle")
        pstyle.set(qn("w:val"), style_name)
        ppr.append(pstyle)
        new_para_element.append(ppr)

    # Add run with text
    if new_text:
        run_element = OxmlElement("w:r")
        text_element = OxmlElement("w:t")
        text_element.text = new_text
        # Preserve leading/trailing spaces
        if new_text != new_text.strip():
            text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        run_element.append(text_element)
        new_para_element.append(run_element)

    # Insert after the reference element in the body
    ref_element.addnext(new_para_element)

    # Return a Paragraph wrapper for the new element
    return Paragraph(new_para_element, doc)


def _delete_paragraph(para: Paragraph) -> None:
    """
    Remove a paragraph from the document body.

    Note: We never delete the very last paragraph in a DOCX body —
    Word requires at least one paragraph. If this is the last one,
    we clear its text instead.
    """
    p_element = para._p
    parent = p_element.getparent()

    if parent is None:
        return

    # Count sibling paragraphs
    siblings = parent.findall(qn("w:p"))
    if len(siblings) <= 1:
        # Last paragraph — clear text instead of removing
        para.clear()
        return

    parent.remove(p_element)


# ── Headline / Summary applier ────────────────────────────────────────────────

def _apply_headline_summary(
    doc: Document,
    new_headline: str | None,
    new_summary: str | None,
) -> None:
    """
    Apply new headline and/or summary to the top of the document.

    Strategy:
    - Headline → replace the FIRST non-empty paragraph text
    - Summary  → replace the SECOND non-empty paragraph text

    This handles the common CV structure:
      [Name / Headline]       ← paragraph 0 or first non-empty
      [Summary paragraph]     ← paragraph 1 or second non-empty

    If the document doesn't have two non-empty paragraphs at the top,
    we apply as much as we can without inserting new paragraphs
    (that would change the layout).
    """
    paragraphs = _get_all_body_paragraphs(doc)

    # Find first two non-empty paragraphs
    non_empty = [p for p in paragraphs if p.text.strip()]

    if new_headline and len(non_empty) >= 1:
        _apply_replace_text(non_empty[0], new_headline)
        logger.debug("render_headline_applied")

    if new_summary and len(non_empty) >= 2:
        _apply_replace_text(non_empty[1], new_summary)
        logger.debug("render_summary_applied")


# ── Main renderer ─────────────────────────────────────────────────────────────

@dataclass
class RenderResult:
    """Result of a render operation."""
    docx_bytes: bytes
    edits_applied: int
    edits_skipped: int
    warnings: list[str]

    def to_metadata(self) -> dict:
        return {
            "edits_applied": self.edits_applied,
            "edits_skipped": self.edits_skipped,
            "warnings": self.warnings,
            "output_size_bytes": len(self.docx_bytes),
        }


class DOCXRenderer:
    """
    Applies an EditPlan to a DOCX file and returns the modified bytes.

    Usage:
        renderer = DOCXRenderer()
        result = renderer.render(docx_bytes=original_bytes, edit_plan=edit_plan_dict)
        # result.docx_bytes is the tailored DOCX
    """

    def render(
        self,
        docx_bytes: bytes,
        edit_plan: dict,
    ) -> RenderResult:
        """
        Apply edit_plan to docx_bytes and return modified DOCX bytes.

        Args:
            docx_bytes: Raw bytes of the original DOCX (from S3)
            edit_plan:  EditPlan dict (as stored in JobRun.edit_plan)

        Returns:
            RenderResult with modified DOCX bytes and operation counts

        Raises:
            ValueError: if docx_bytes is not a valid DOCX
        """
        # Load document
        try:
            doc = Document(io.BytesIO(docx_bytes))
        except Exception as e:
            raise ValueError(f"Could not open DOCX for rendering: {e}") from e

        edits_applied = 0
        edits_skipped = 0
        warnings: list[str] = []

        # ── Step 1: Headline / Summary ─────────────────────────────────────
        headline_edit = edit_plan.get("headline_summary")
        if headline_edit:
            try:
                _apply_headline_summary(
                    doc=doc,
                    new_headline=headline_edit.get("new_headline"),
                    new_summary=headline_edit.get("new_summary"),
                )
                edits_applied += 1
            except Exception as e:
                warnings.append(f"headline_summary failed: {e}")
                edits_skipped += 1

        # ── Step 2: Paragraph edits ────────────────────────────────────────
        # Sort edits: process INSERT and DELETE in reverse index order
        # so that earlier indices remain valid as we modify the document.
        # REPLACE_TEXT and REPLACE_RUN can be applied in any order.
        paragraph_edits = edit_plan.get("paragraph_edits", [])

        # Separate by operation type
        insert_delete_edits = [
            e for e in paragraph_edits
            if e.get("operation") in ("insert_after", "delete")
        ]
        replace_edits = [
            e for e in paragraph_edits
            if e.get("operation") in ("replace_text", "replace_run", "restyle")
        ]

        # Apply replaces first (index-stable)
        for edit in replace_edits:
            applied, warning = self._apply_paragraph_edit(doc, edit)
            if applied:
                edits_applied += 1
            else:
                edits_skipped += 1
                if warning:
                    warnings.append(warning)

        # Apply inserts/deletes in reverse order (index-shifting)
        for edit in sorted(
            insert_delete_edits,
            key=lambda e: e.get("paragraph_index", 0),
            reverse=True,
        ):
            applied, warning = self._apply_paragraph_edit(doc, edit)
            if applied:
                edits_applied += 1
            else:
                edits_skipped += 1
                if warning:
                    warnings.append(warning)

        # ── Step 3: Sections to add (append at end) ────────────────────────
        for section in edit_plan.get("sections_to_add", []):
            try:
                heading = section.get("heading", "")
                content = section.get("content", "")
                if heading:
                    doc.add_paragraph(heading, style="Heading 1")
                    edits_applied += 1
                if content:
                    doc.add_paragraph(content)
                    edits_applied += 1
            except Exception as e:
                warnings.append(f"sections_to_add failed: {e}")
                edits_skipped += 1

        logger.info(
            "render_complete",
            extra={
                "edits_applied": edits_applied,
                "edits_skipped": edits_skipped,
                "warnings": len(warnings),
            },
        )

        # ── Step 4: Serialise to bytes ─────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes_out = buf.getvalue()

        return RenderResult(
            docx_bytes=docx_bytes_out,
            edits_applied=edits_applied,
            edits_skipped=edits_skipped,
            warnings=warnings,
        )

    def _apply_paragraph_edit(
        self,
        doc: Document,
        edit: dict,
    ) -> tuple[bool, str | None]:
        """
        Apply a single paragraph edit operation.

        Returns:
            (success: bool, warning_message: str | None)
        """
        paragraphs = _get_all_body_paragraphs(doc)
        idx = edit.get("paragraph_index", -1)
        operation = edit.get("operation", "")
        new_text = edit.get("new_text") or ""
        run_index = edit.get("run_index")
        style = edit.get("style")

        # Bounds check for operations that require existing paragraph
        if operation != "insert_after" and (idx < 0 or idx >= len(paragraphs)):
            return False, (
                f"paragraph_index {idx} out of range "
                f"(document has {len(paragraphs)} paragraphs)"
            )

        try:
            if operation == "replace_text":
                _apply_replace_text(paragraphs[idx], new_text)
                return True, None

            elif operation == "replace_run":
                if run_index is None:
                    return False, f"replace_run at index {idx} missing run_index"
                _apply_replace_run(paragraphs[idx], run_index, new_text)
                return True, None

            elif operation == "insert_after":
                _insert_paragraph_after(
                    doc=doc,
                    after_index=idx,
                    new_text=new_text,
                    style_name=style,
                )
                return True, None

            elif operation == "delete":
                _delete_paragraph(paragraphs[idx])
                return True, None

            elif operation == "restyle":
                if not style:
                    return False, f"restyle at index {idx} missing style name"
                _apply_restyle(paragraphs[idx], style, doc)
                return True, None

            else:
                return False, f"Unknown operation: {operation}"

        except Exception as e:
            return False, f"Edit {operation}@{idx} failed: {e}"
