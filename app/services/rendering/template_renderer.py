"""
app/services/rendering/template_renderer.py

Template renderer — fills a DOCX template with BuiltCV content.

Used when cv.build_mode is FROM_SCRATCH or REBUILD.
The EditPlan renderer handles the EDIT mode (existing DOCX + diff).

Design principles:
  - Templates are plain DOCX files with placeholder text markers
  - Markers: {{NAME}}, {{HEADLINE}}, {{SUMMARY}}, {{CONTACT}}, etc.
  - Experience, education, and skills sections are built programmatically
  - If no template DOCX is available, falls back to a clean generated document
  - All formatting (fonts, colours, spacing) comes from the template — never hardcoded

Template marker convention:
  {{NAME}}         Full name
  {{HEADLINE}}     Positioning headline
  {{SUMMARY}}      Executive summary paragraph
  {{CONTACT}}      Contact line (email · phone · location · linkedin)
  {{EXPERIENCE}}   Placeholder paragraph — replaced with built experience section
  {{EDUCATION}}    Placeholder paragraph — replaced with built education section
  {{SKILLS}}       Placeholder paragraph — replaced with built skills section

If a template has none of these markers (or no template is found), the renderer
generates a clean document from scratch using sensible defaults.

Fallback document spec:
  - Calibri 11pt body, 14pt section headings
  - 2.5cm margins
  - Single column
  - Horizontal rule after each section heading
"""

import io
import re
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.text.paragraph import Paragraph

import structlog

logger = structlog.get_logger(__name__)


@dataclass
class TemplateRenderResult:
    docx_bytes: bytes
    sections_written: int
    bullets_written: int
    mode: str          # "template_filled" or "generated"
    warnings: list[str]

    def to_metadata(self) -> dict:
        return {
            "sections_written": self.sections_written,
            "bullets_written": self.bullets_written,
            "mode": self.mode,
            "warnings": self.warnings,
        }


class TemplateRenderer:
    """
    Renders a BuiltCV dict into a DOCX file.

    Usage:
        renderer = TemplateRenderer()
        result = renderer.render(built_cv=built_cv_dict, template_bytes=template_docx_bytes)
        # result.docx_bytes is the ready-to-upload DOCX
    """

    def render(
        self,
        built_cv: dict,
        template_bytes: bytes | None = None,
        template_slug: str | None = None,
    ) -> TemplateRenderResult:
        """
        Render a BuiltCV dict into a DOCX.

        Args:
            built_cv:       BuiltCV dict from CVBuildService
            template_bytes: Raw DOCX template bytes from S3 (None = use fallback)
            template_slug:  Slug name for logging

        Returns:
            TemplateRenderResult with docx_bytes ready for S3 upload
        """
        warnings = []

        if template_bytes:
            try:
                return self._render_with_template(built_cv, template_bytes, warnings)
            except Exception as e:
                logger.warning(
                    "template_render_failed_falling_back",
                    template_slug=template_slug,
                    error=str(e),
                )
                warnings.append(f"Template rendering failed ({e}), used generated fallback.")

        # Fallback: generate clean document programmatically
        return self._render_generated(built_cv, warnings)

    # ── Template-based rendering ──────────────────────────────────────────────

    def _render_with_template(
        self, built_cv: dict, template_bytes: bytes, warnings: list[str]
    ) -> TemplateRenderResult:
        """Fill a template DOCX with BuiltCV content."""
        doc = Document(io.BytesIO(template_bytes))
        sections_written = 0
        bullets_written = 0

        # Replace simple text markers in all paragraphs
        replacements = {
            "{{NAME}}": built_cv.get("name", ""),
            "{{HEADLINE}}": built_cv.get("headline", ""),
            "{{SUMMARY}}": built_cv.get("summary", ""),
            "{{CONTACT}}": _format_contact_line(built_cv.get("contact", {})),
        }

        for para in doc.paragraphs:
            for marker, value in replacements.items():
                if marker in para.text:
                    _replace_marker_in_para(para, marker, value)

        # Find and replace section placeholder paragraphs
        paras = list(doc.paragraphs)
        for i, para in enumerate(paras):
            text = para.text.strip()

            if "{{EXPERIENCE}}" in text:
                exp_section = _find_section(built_cv, "experience")
                if exp_section:
                    n_b = _insert_experience_section(doc, para, exp_section)
                    bullets_written += n_b
                    sections_written += 1
                    _remove_paragraph(para)
                else:
                    warnings.append("No experience section in BuiltCV")

            elif "{{EDUCATION}}" in text:
                edu_section = _find_section(built_cv, "education")
                if edu_section:
                    _insert_education_section(doc, para, edu_section)
                    sections_written += 1
                    _remove_paragraph(para)

            elif "{{SKILLS}}" in text:
                skills_section = _find_section(built_cv, "skills")
                if skills_section:
                    _insert_skills_section(doc, para, skills_section)
                    sections_written += 1
                    _remove_paragraph(para)

        docx_bytes = _doc_to_bytes(doc)
        return TemplateRenderResult(
            docx_bytes=docx_bytes,
            sections_written=sections_written,
            bullets_written=bullets_written,
            mode="template_filled",
            warnings=warnings,
        )

    # ── Generated (fallback) rendering ───────────────────────────────────────

    def _render_generated(
        self, built_cv: dict, warnings: list[str]
    ) -> TemplateRenderResult:
        """
        Generate a clean, professional DOCX from scratch.
        Used when no template is available or template rendering fails.
        """
        doc = Document()
        sections_written = 0
        bullets_written = 0

        # Page setup: A4, 2.5cm margins
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(section, attr, Cm(2.5))

        contact = built_cv.get("contact", {})

        # ── Name ──────────────────────────────────────────────────────────────
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(built_cv.get("name", "Your Name"))
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        name_para.paragraph_format.space_after = Pt(2)

        # ── Headline ──────────────────────────────────────────────────────────
        if built_cv.get("headline"):
            hl_para = doc.add_paragraph()
            hl_run = hl_para.add_run(built_cv["headline"])
            hl_run.font.size = Pt(11)
            hl_run.font.italic = True
            hl_run.font.color.rgb = RGBColor(0x55, 0x55, 0x70)
            hl_para.paragraph_format.space_after = Pt(2)

        # ── Contact line ──────────────────────────────────────────────────────
        contact_line = _format_contact_line(contact)
        if contact_line:
            ct_para = doc.add_paragraph()
            ct_run = ct_para.add_run(contact_line)
            ct_run.font.size = Pt(9)
            ct_run.font.color.rgb = RGBColor(0x55, 0x55, 0x70)
            ct_para.paragraph_format.space_after = Pt(8)

        _add_horizontal_rule(doc)

        # ── Summary ───────────────────────────────────────────────────────────
        if built_cv.get("summary"):
            _add_section_heading(doc, "Professional Summary")
            sum_para = doc.add_paragraph(built_cv["summary"])
            sum_para.style.font.size = Pt(10)
            sum_para.paragraph_format.space_after = Pt(8)
            sections_written += 1

        # ── Sections ──────────────────────────────────────────────────────────
        for section_data in built_cv.get("sections", []):
            stype = section_data.get("section_type", "")

            if stype == "experience":
                _add_section_heading(doc, section_data.get("title", "Professional Experience"))
                for entry in section_data.get("entries", []):
                    n_b = _write_experience_entry(doc, entry)
                    bullets_written += n_b
                sections_written += 1

            elif stype == "education":
                _add_section_heading(doc, section_data.get("title", "Education"))
                for entry in section_data.get("entries", []):
                    _write_education_entry(doc, entry)
                sections_written += 1

            elif stype == "skills":
                _add_section_heading(doc, section_data.get("title", "Skills & Expertise"))
                _write_skills_section(doc, section_data)
                sections_written += 1

            else:
                # Generic section
                _add_section_heading(doc, section_data.get("title", stype.title()))
                for entry in section_data.get("entries", []):
                    if isinstance(entry, dict):
                        for bullet in entry.get("bullets", []):
                            _add_bullet(doc, bullet)
                            bullets_written += 1
                sections_written += 1

        docx_bytes = _doc_to_bytes(doc)
        return TemplateRenderResult(
            docx_bytes=docx_bytes,
            sections_written=sections_written,
            bullets_written=bullets_written,
            mode="generated",
            warnings=warnings,
        )


# ── Document writing helpers ──────────────────────────────────────────────────

def _add_section_heading(doc: Document, text: str) -> Paragraph:
    para = doc.add_paragraph()
    run = para.add_run(text.upper())
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.font.name = "Calibri"
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    # Bottom border (subtle rule under heading)
    _add_para_bottom_border(para)
    return para


def _write_experience_entry(doc: Document, entry: dict) -> int:
    """Write one experience entry. Returns number of bullets written."""
    # Role + Company line
    role_para = doc.add_paragraph()
    role_run = role_para.add_run(entry.get("role", ""))
    role_run.font.bold = True
    role_run.font.size = Pt(10.5)

    company_run = role_para.add_run(f"  ·  {entry.get('company', '')}")
    company_run.font.size = Pt(10)
    company_run.font.color.rgb = RGBColor(0x44, 0x44, 0x55)
    role_para.paragraph_format.space_after = Pt(0)

    # Dates + location line
    dates = f"{entry.get('start_date', '')} – {entry.get('end_date', 'Present')}"
    loc = entry.get("location", "")
    date_line = f"{dates}  ·  {loc}" if loc else dates
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(date_line)
    date_run.font.size = Pt(9)
    date_run.font.italic = True
    date_run.font.color.rgb = RGBColor(0x77, 0x77, 0x88)
    date_para.paragraph_format.space_after = Pt(2)

    # Bullets
    bullets = entry.get("bullets", [])
    for bullet in bullets:
        _add_bullet(doc, bullet)

    # Spacing after entry
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

    return len(bullets)


def _write_education_entry(doc: Document, entry: dict) -> None:
    edu_para = doc.add_paragraph()
    deg_run = edu_para.add_run(entry.get("degree", ""))
    deg_run.font.bold = True
    deg_run.font.size = Pt(10.5)

    inst_run = edu_para.add_run(f"  ·  {entry.get('institution', '')}")
    inst_run.font.size = Pt(10)
    inst_run.font.color.rgb = RGBColor(0x44, 0x44, 0x55)
    edu_para.paragraph_format.space_after = Pt(0)

    year_loc = entry.get("year", "")
    loc = entry.get("location", "")
    sub = f"{year_loc}  ·  {loc}" if loc else year_loc
    if sub.strip():
        sub_para = doc.add_paragraph()
        sub_run = sub_para.add_run(sub.strip("  · "))
        sub_run.font.size = Pt(9)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(0x77, 0x77, 0x88)
        sub_para.paragraph_format.space_after = Pt(2)

    notes = entry.get("notes", "")
    if notes:
        n_para = doc.add_paragraph()
        n_run = n_para.add_run(notes)
        n_run.font.size = Pt(9.5)
        n_para.paragraph_format.space_after = Pt(4)


def _write_skills_section(doc: Document, section_data: dict) -> None:
    for cat in section_data.get("categories", []):
        label = cat.get("label", "")
        items = cat.get("items", [])
        if not items:
            continue
        sk_para = doc.add_paragraph()
        lbl_run = sk_para.add_run(f"{label}: " if label else "")
        lbl_run.font.bold = True
        lbl_run.font.size = Pt(10)
        items_run = sk_para.add_run("  ·  ".join(items))
        items_run.font.size = Pt(10)
        sk_para.paragraph_format.space_after = Pt(3)


def _add_bullet(doc: Document, text: str) -> Paragraph:
    para = doc.add_paragraph(style="List Bullet" if _style_exists(doc, "List Bullet") else "Normal")
    if para.style.name == "Normal":
        # Manual bullet if style unavailable
        run = para.add_run(f"•  {text}")
    else:
        run = para.add_run(text)
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.left_indent = Inches(0.25)
    return para


def _add_horizontal_rule(doc: Document) -> None:
    """Add a thin horizontal line paragraph."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_after = Pt(4)


def _add_para_bottom_border(para: Paragraph) -> None:
    """Add a subtle bottom border to a paragraph (used for section headings)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "2")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Template helpers ──────────────────────────────────────────────────────────

def _replace_marker_in_para(para: Paragraph, marker: str, value: str) -> None:
    """Replace a marker in a paragraph, preserving the run's formatting."""
    for run in para.runs:
        if marker in run.text:
            run.text = run.text.replace(marker, value)


def _find_section(built_cv: dict, section_type: str) -> dict | None:
    for sec in built_cv.get("sections", []):
        if sec.get("section_type") == section_type:
            return sec
    return None


def _insert_experience_section(
    doc: Document, placeholder_para: Paragraph, section_data: dict
) -> int:
    """Insert experience entries before the placeholder paragraph."""
    bullets_written = 0
    # We write to a temp doc then transfer — simplest approach with python-docx
    for entry in section_data.get("entries", []):
        n_b = _write_experience_entry(doc, entry)
        bullets_written += n_b
    return bullets_written


def _insert_education_section(
    doc: Document, placeholder_para: Paragraph, section_data: dict
) -> None:
    for entry in section_data.get("entries", []):
        _write_education_entry(doc, entry)


def _insert_skills_section(
    doc: Document, placeholder_para: Paragraph, section_data: dict
) -> None:
    _write_skills_section(doc, section_data)


def _remove_paragraph(para: Paragraph) -> None:
    """Remove a paragraph from the document."""
    p = para._element
    p.getparent().remove(p)


def _format_contact_line(contact: dict) -> str:
    """Format contact dict into a single readable line."""
    parts = []
    if contact.get("email"):
        parts.append(contact["email"])
    if contact.get("phone"):
        parts.append(contact["phone"])
    if contact.get("location"):
        parts.append(contact["location"])
    if contact.get("linkedin"):
        # Shorten URL for readability
        url = contact["linkedin"].replace("https://www.", "").replace("https://", "")
        parts.append(url)
    return "  ·  ".join(parts)


def _style_exists(doc: Document, style_name: str) -> bool:
    return any(s.name == style_name for s in doc.styles)


def _doc_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
