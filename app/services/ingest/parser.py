"""
app/services/ingest/parser.py

DOCX parser and node mapper.

Converts a raw DOCX binary into a structured ParsedCV object.
This is the data the LLM will read — so fidelity matters.

What we extract per paragraph:
  - index:    position in document (0-based)
  - text:     full plain text
  - style:    paragraph style name (Heading 1, Normal, etc.)
  - runs:     individual text runs with formatting metadata
  - is_empty: True for whitespace-only paragraphs

Section detection:
  Paragraphs with Heading styles are treated as section boundaries.
  Everything between headings is grouped into that section.
  If no headings found, all paragraphs go into a single 'Document' section.

Design constraint: This module is synchronous — it runs inside a Celery
worker. Do not use async/await here.
"""

import io
import re
from dataclasses import dataclass, field

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from app.schemas.cv import ParsedCV, ParsedNode, ParsedSection

import structlog

logger = structlog.get_logger(__name__)

# ── Heading detection ──────────────────────────────────────────────────────
# Paragraph styles that we treat as section boundaries
HEADING_STYLE_PREFIXES = ("heading", "title")

# All-caps paragraphs that are short enough to be section headers
# (common in CVs that don't use Word heading styles)
ALL_CAPS_MAX_WORDS = 6


def _is_heading(para) -> bool:
    """
    Detect whether a python-docx paragraph is a section heading.

    Two signals:
    1. Word style name starts with 'Heading' or is 'Title'
    2. Short paragraph in ALL CAPS (common CV formatting)
    """
    style_name = (para.style.name or "").lower()
    if any(style_name.startswith(p) for p in HEADING_STYLE_PREFIXES):
        return True

    text = para.text.strip()
    if not text:
        return False

    words = text.split()
    if len(words) <= ALL_CAPS_MAX_WORDS and text == text.upper() and len(text) > 2:
        return True

    return False


def _extract_run_metadata(run) -> dict:
    """
    Extract formatting metadata from a single python-docx Run.

    We capture enough to reconstruct the run's appearance when
    applying edits — but we don't need everything (e.g. colour, spacing).
    """
    font = run.font

    # Font size: run-level first, fall back to paragraph style
    font_size: float | None = None
    if font.size is not None:
        font_size = font.size.pt  # Convert EMUs → points

    return {
        "text": run.text,
        "bold": run.bold or False,
        "italic": run.italic or False,
        "underline": run.underline or False,
        "font_name": font.name,
        "font_size_pt": font_size,
    }


def _count_words(text: str) -> int:
    """Count non-empty whitespace-separated tokens."""
    return len(text.split())


# ── Main parser ────────────────────────────────────────────────────────────

class DOCXParser:
    """
    Parses a DOCX binary into a structured ParsedCV.

    Usage:
        parser = DOCXParser()
        parsed = parser.parse(docx_bytes)
    """

    def parse(self, docx_bytes: bytes) -> ParsedCV:
        """
        Parse raw DOCX bytes into a ParsedCV object.

        Args:
            docx_bytes: Raw bytes of a DOCX file (from S3)

        Returns:
            ParsedCV with sections, paragraphs, and word count

        Raises:
            ValueError: if the bytes are not a valid DOCX file
        """
        try:
            doc = Document(io.BytesIO(docx_bytes))
        except Exception as e:
            raise ValueError(f"Could not open DOCX file: {e}") from e

        # ── Collect ALL paragraphs including those inside tables ───────────
        # doc.paragraphs only returns top-level paragraphs — misses table cells.
        # Many professional CVs use table layouts, so we must walk the XML.
        from docx.oxml.ns import qn as _qn
        from lxml import etree as _etree

        all_paras = []
        body = doc.element.body
        for child in body.iter():
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'p':  # paragraph element
                from docx.text.paragraph import Paragraph as _Para
                try:
                    p = _Para(child, doc)
                    all_paras.append(p)
                except Exception:
                    pass

        # Deduplicate by object id — doc.paragraphs are already in all_paras
        seen_ids = set()
        unique_paras = []
        for p in all_paras:
            pid = id(p._element)
            if pid not in seen_ids:
                seen_ids.add(pid)
                unique_paras.append(p)

        paragraphs = unique_paras
        total_paragraphs = len(paragraphs)

        logger.info(
            "docx_parse_start",
            extra={"total_paragraphs": total_paragraphs},
        )

        # ── Build flat node list ───────────────────────────────────────────
        nodes: list[ParsedNode] = []
        total_words = 0

        for idx, para in enumerate(paragraphs):
            text = para.text  # Full text, runs concatenated
            stripped = text.strip()
            is_empty = len(stripped) == 0

            runs = [_extract_run_metadata(r) for r in para.runs if r.text]

            node = ParsedNode(
                index=idx,
                text=text,
                style=para.style.name or "Normal",
                runs=runs,
                is_empty=is_empty,
            )
            nodes.append(node)

            if not is_empty:
                total_words += _count_words(stripped)

        # ── Group into sections ────────────────────────────────────────────
        sections = self._group_into_sections(nodes)

        # ── Fallback: if no sections detected, return single section ──────
        if not sections:
            sections = [
                ParsedSection(
                    heading="Document",
                    heading_index=0,
                    paragraphs=nodes,
                )
            ]

        logger.info(
            "docx_parse_complete",
            extra={
                "sections": len(sections),
                "paragraphs": total_paragraphs,
                "words": total_words,
            },
        )

        return ParsedCV(
            total_paragraphs=total_paragraphs,
            total_words=total_words,
            sections=sections,
            raw_paragraphs=nodes,
        )

    def _group_into_sections(
        self,
        nodes: list[ParsedNode],
    ) -> list[ParsedSection]:
        """
        Group flat paragraph nodes into sections using heading detection.

        Algorithm:
        - Walk paragraphs in order
        - When a heading is found, start a new section
        - Paragraphs before any heading go into a 'Preamble' section
        - Empty paragraphs are included (they carry formatting intent)
        """
        sections: list[ParsedSection] = []
        current_heading = "Preamble"
        current_heading_index = 0
        current_paragraphs: list[ParsedNode] = []

        for node in nodes:
            # We need the original para to check style — reconstruct from node
            # heading detection is based on style name and text
            if self._node_is_heading(node):
                # Save current section if it has content
                if current_paragraphs:
                    sections.append(
                        ParsedSection(
                            heading=current_heading,
                            heading_index=current_heading_index,
                            paragraphs=current_paragraphs,
                        )
                    )
                # Start new section
                current_heading = node.text.strip() or "Section"
                current_heading_index = node.index
                current_paragraphs = [node]  # Include heading node itself
            else:
                current_paragraphs.append(node)

        # Don't forget the last section
        if current_paragraphs:
            sections.append(
                ParsedSection(
                    heading=current_heading,
                    heading_index=current_heading_index,
                    paragraphs=current_paragraphs,
                )
            )

        return sections

    def _node_is_heading(self, node: ParsedNode) -> bool:
        """
        Heading detection on a ParsedNode (post-parse, no access to docx para).
        Uses style name and ALL_CAPS heuristic.
        """
        style_lower = node.style.lower()
        if any(style_lower.startswith(p) for p in HEADING_STYLE_PREFIXES):
            return True

        text = node.text.strip()
        if not text or node.is_empty:
            return False

        words = text.split()
        if (
            len(words) <= ALL_CAPS_MAX_WORDS
            and text == text.upper()
            and len(text) > 2
        ):
            return True

        return False
