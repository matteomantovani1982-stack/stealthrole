"""
tests/test_parser.py

Tests for the DOCX parser and parse_cv Celery task.

Parser tests use real python-docx objects built in memory — no files on disk.
Task tests mock S3 and DB — no infrastructure needed.
"""

import io
import uuid
from unittest.mock import MagicMock, patch, call

import pytest
from docx import Document
from docx.shared import Pt

from app.services.ingest.parser import DOCXParser, _is_heading, _extract_run_metadata
from app.schemas.cv import ParsedCV, ParsedSection, ParsedNode


# ── Helpers: build test DOCX documents in memory ───────────────────────────

def make_docx_bytes(paragraphs: list[tuple[str, str]]) -> bytes:
    """
    Build a DOCX in memory with given (text, style_name) paragraphs.
    Returns raw bytes as if downloaded from S3.

    Example:
        make_docx_bytes([
            ("Matteo Mantovani", "Heading 1"),
            ("Senior Strategy Executive", "Normal"),
            ("EXPERIENCE", "Normal"),        # ALL CAPS heading
            ("McKinsey & Company", "Normal"),
        ])
    """
    doc = Document()
    for text, style in paragraphs:
        try:
            para = doc.add_paragraph(text, style=style)
        except KeyError:
            # Style doesn't exist in this template — use Normal
            para = doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_simple_cv_bytes() -> bytes:
    """A realistic CV structure for integration-style tests."""
    return make_docx_bytes([
        ("Matteo Mantovani", "Normal"),
        ("CEO | Tech Founder | INSEAD MBA", "Normal"),
        ("Dubai, UAE  |  matteo@example.com", "Normal"),
        ("", "Normal"),
        ("EXPERIENCE", "Normal"),               # ALL CAPS section
        ("CEO, Liively — Dubai (2024–present)", "Normal"),
        ("Built product from 0 to $1M ARR in 12 months.", "Normal"),
        ("", "Normal"),
        ("CEO, LaLiga Iraqi Stars League (2023–2024)", "Normal"),
        ("Transformed struggling league into viable business.", "Normal"),
        ("", "Normal"),
        ("EDUCATION", "Normal"),                # ALL CAPS section
        ("INSEAD — MBA (2018–2019)", "Normal"),
        ("Imperial College London — MSc Materials Engineering", "Normal"),
    ])


# ── Unit tests: heading detection ──────────────────────────────────────────

class TestHeadingDetection:
    """Tests for the _node_is_heading heuristic on ParsedNode objects."""

    def _make_node(self, text: str, style: str = "Normal") -> ParsedNode:
        return ParsedNode(
            index=0,
            text=text,
            style=style,
            runs=[],
            is_empty=(len(text.strip()) == 0),
        )

    def test_heading1_style_detected(self):
        node = self._make_node("Work Experience", style="Heading 1")
        parser = DOCXParser()
        assert parser._node_is_heading(node) is True

    def test_heading2_style_detected(self):
        node = self._make_node("McKinsey & Company", style="Heading 2")
        parser = DOCXParser()
        assert parser._node_is_heading(node) is True

    def test_all_caps_short_detected(self):
        node = self._make_node("EXPERIENCE")
        parser = DOCXParser()
        assert parser._node_is_heading(node) is True

    def test_all_caps_education_detected(self):
        node = self._make_node("EDUCATION")
        parser = DOCXParser()
        assert parser._node_is_heading(node) is True

    def test_all_caps_long_not_detected(self):
        # 7 words — over ALL_CAPS_MAX_WORDS threshold
        node = self._make_node("THIS IS A VERY LONG ALL CAPS SENTENCE NOT HEADING")
        parser = DOCXParser()
        assert parser._node_is_heading(node) is False

    def test_normal_sentence_not_detected(self):
        node = self._make_node("Built Iraq's first super-app from zero to 500 employees.")
        parser = DOCXParser()
        assert parser._node_is_heading(node) is False

    def test_empty_not_detected(self):
        node = self._make_node("")
        parser = DOCXParser()
        assert parser._node_is_heading(node) is False

    def test_mixed_case_not_detected(self):
        node = self._make_node("McKinsey & Company")
        parser = DOCXParser()
        assert parser._node_is_heading(node) is False

    def test_two_char_all_caps_not_detected(self):
        # Too short to be a section heading
        node = self._make_node("UK")
        parser = DOCXParser()
        assert parser._node_is_heading(node) is False


# ── Unit tests: run extraction ─────────────────────────────────────────────

class TestRunExtraction:
    """Tests for run-level formatting extraction."""

    def test_basic_run_extracted(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Hello World")
        run.bold = True

        result = _extract_run_metadata(run)

        assert result["text"] == "Hello World"
        assert result["bold"] is True
        assert result["italic"] is False

    def test_italic_run_extracted(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Italic text")
        run.italic = True

        result = _extract_run_metadata(run)
        assert result["italic"] is True
        assert result["bold"] is False

    def test_font_size_extracted(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Large text")
        run.font.size = Pt(14)

        result = _extract_run_metadata(run)
        assert result["font_size_pt"] == 14.0


# ── Integration tests: full DOCX parse ─────────────────────────────────────

class TestDOCXParser:
    """Integration tests parsing real in-memory DOCX files."""

    def test_parse_returns_parsed_cv(self):
        docx_bytes = make_simple_cv_bytes()
        parser = DOCXParser()
        result = parser.parse(docx_bytes)

        assert isinstance(result, ParsedCV)
        assert result.total_paragraphs > 0
        assert result.total_words > 0

    def test_all_caps_sections_detected(self):
        docx_bytes = make_simple_cv_bytes()
        parser = DOCXParser()
        result = parser.parse(docx_bytes)

        section_headings = [s.heading for s in result.sections]
        assert "EXPERIENCE" in section_headings
        assert "EDUCATION" in section_headings

    def test_paragraphs_have_correct_text(self):
        docx_bytes = make_simple_cv_bytes()
        parser = DOCXParser()
        result = parser.parse(docx_bytes)

        all_texts = [p.text for s in result.sections for p in s.paragraphs]
        assert any("Liively" in t for t in all_texts)
        assert any("INSEAD" in t for t in all_texts)

    def test_empty_paragraphs_preserved(self):
        """Empty paragraphs must be preserved — they carry formatting intent."""
        docx_bytes = make_simple_cv_bytes()
        parser = DOCXParser()
        result = parser.parse(docx_bytes)

        all_nodes = result.raw_paragraphs
        empty_nodes = [n for n in all_nodes if n.is_empty]
        assert len(empty_nodes) > 0

    def test_word_count_is_nonzero(self):
        docx_bytes = make_simple_cv_bytes()
        parser = DOCXParser()
        result = parser.parse(docx_bytes)
        assert result.total_words > 50

    def test_raw_paragraphs_count_matches_total(self):
        docx_bytes = make_simple_cv_bytes()
        parser = DOCXParser()
        result = parser.parse(docx_bytes)
        assert len(result.raw_paragraphs) == result.total_paragraphs

    def test_invalid_bytes_raises_value_error(self):
        with pytest.raises(ValueError, match="Could not open DOCX"):
            DOCXParser().parse(b"this is not a docx file at all")

    def test_empty_document_returns_single_section(self):
        """Document with no recognisable headings gets a single fallback section."""
        docx_bytes = make_docx_bytes([
            ("Just some plain text.", "Normal"),
            ("No headings anywhere.", "Normal"),
        ])
        parser = DOCXParser()
        result = parser.parse(docx_bytes)
        assert len(result.sections) >= 1

    def test_heading1_style_creates_sections(self):
        """Heading 1 style paragraphs should create section boundaries."""
        docx_bytes = make_docx_bytes([
            ("Work Experience", "Heading 1"),
            ("McKinsey & Company", "Normal"),
            ("Education", "Heading 1"),
            ("INSEAD MBA", "Normal"),
        ])
        parser = DOCXParser()
        result = parser.parse(docx_bytes)

        headings = [s.heading for s in result.sections]
        assert "Work Experience" in headings
        assert "Education" in headings

    def test_parsed_cv_serialises_to_dict(self):
        """parsed_content is stored as dict in DB — must be serialisable."""
        docx_bytes = make_simple_cv_bytes()
        parser = DOCXParser()
        result = parser.parse(docx_bytes)

        data = result.model_dump()
        assert isinstance(data, dict)
        assert "sections" in data
        assert "total_words" in data
        assert "raw_paragraphs" in data


# ── Task tests ─────────────────────────────────────────────────────────────

class TestParseCVTask:
    """
    Tests for the parse_cv_task Celery task.

    All external dependencies (S3, DB) are mocked.
    We test the task's logic and state transitions.
    """

    def _make_mock_cv(self, status="uploaded"):
        from app.models.cv import CVStatus
        cv = MagicMock()
        cv.id = uuid.uuid4()
        cv.s3_key = f"cvs/user/{cv.id}/cv.docx"
        cv.s3_bucket = "careeros"
        cv.status = CVStatus.UPLOADED
        cv.parsed_content = None
        cv.error_message = None
        return cv

    @patch("app.workers.tasks.parse_cv.get_sync_db")
    @patch("app.services.ingest.storage.S3StorageService.download_bytes")
    @patch("app.services.ingest.parser.DOCXParser.parse")
    def test_successful_parse_updates_status_to_parsed(
        self,
        mock_parse,
        mock_download,
        mock_get_db,
    ):
        """Happy path: task should set status=PARSED and store parsed_content."""
        from app.models.cv import CVStatus
        from app.schemas.cv import ParsedCV

        cv_id = uuid.uuid4()
        mock_cv = self._make_mock_cv()
        mock_cv.id = cv_id

        # Mock DB context manager
        mock_db = MagicMock()
        mock_db.get.return_value = mock_cv
        mock_get_db.return_value.__enter__ = MagicMock(return_value=mock_db)
        mock_get_db.return_value.__exit__ = MagicMock(return_value=False)

        # Mock S3 download
        mock_download.return_value = b"fake docx bytes"

        # Mock parser output
        fake_parsed = ParsedCV(
            total_paragraphs=10,
            total_words=500,
            sections=[],
            raw_paragraphs=[],
        )
        mock_parse.return_value = fake_parsed

        from app.workers.tasks.parse_cv import parse_cv_task
        result = parse_cv_task(str(cv_id))

        assert result["status"] == "parsed"
        assert result["total_words"] == 500

    @patch("app.workers.tasks.parse_cv._mark_cv_failed")
    @patch("app.workers.tasks.parse_cv.get_sync_db")
    @patch("app.services.ingest.storage.S3StorageService.download_bytes")
    def test_invalid_docx_marks_cv_failed(
        self,
        mock_download,
        mock_get_db,
        mock_mark_failed,
    ):
        """Invalid DOCX bytes should mark CV as FAILED without retrying."""
        cv_id = uuid.uuid4()
        mock_cv = self._make_mock_cv()

        mock_db = MagicMock()
        mock_db.get.return_value = mock_cv
        mock_get_db.return_value.__enter__ = MagicMock(return_value=mock_db)
        mock_get_db.return_value.__exit__ = MagicMock(return_value=False)

        mock_download.return_value = b"not a docx"

        from app.workers.tasks.parse_cv import parse_cv_task
        with pytest.raises(ValueError):
            parse_cv_task(str(cv_id))

        mock_mark_failed.assert_called_once()

    def test_invalid_uuid_raises_immediately(self):
        """A malformed cv_id should raise ValueError without touching DB."""
        from app.workers.tasks.parse_cv import parse_cv_task
        with pytest.raises(ValueError):
            parse_cv_task("not-a-valid-uuid")
