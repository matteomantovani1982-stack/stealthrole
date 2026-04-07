"""
tests/test_renderer.py

Tests for the DOCX renderer.

Uses real in-memory DOCX files built with python-docx.
No external services needed — these are pure unit/integration tests.

Coverage:
  - All 5 edit operations
  - Headline/summary replacement
  - Edge cases: out-of-bounds indices, missing runs, last paragraph
  - Formatting preservation (run properties survive replacement)
  - Full round-trip: build DOCX → apply EditPlan → re-parse → verify
  - render_docx task (mocked S3 + DB)
"""

import io
import json
import uuid
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.shared import Pt

from app.services.rendering.docx_renderer import (
    DOCXRenderer,
    RenderResult,
    _apply_replace_text,
    _apply_replace_run,
    _apply_restyle,
    _apply_headline_summary,
    _copy_run_format,
    _delete_paragraph,
    _get_all_body_paragraphs,
    _insert_paragraph_after,
)


# ── DOCX builder helpers ─────────────────────────────────────────────────────

def make_docx_bytes(*paragraphs: str | tuple) -> bytes:
    """
    Build an in-memory DOCX with the given paragraphs.

    Each arg can be:
      - str: plain paragraph with default style
      - (text, style): paragraph with explicit style
      - (text, style, bold): paragraph with first run bolded
    """
    doc = Document()
    for p in paragraphs:
        if isinstance(p, str):
            doc.add_paragraph(p)
        elif len(p) == 2:
            text, style = p
            try:
                doc.add_paragraph(text, style=style)
            except KeyError:
                doc.add_paragraph(text)
        elif len(p) == 3:
            text, style, bold = p
            try:
                para = doc.add_paragraph(style=style)
            except KeyError:
                para = doc.add_paragraph()
            run = para.add_run(text)
            run.bold = bold
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_paragraphs(docx_bytes: bytes) -> list[str]:
    """Extract all paragraph texts from DOCX bytes."""
    doc = Document(io.BytesIO(docx_bytes))
    return [p.text for p in doc.paragraphs]


def docx_run_is_bold(docx_bytes: bytes, para_idx: int, run_idx: int = 0) -> bool:
    """Check if a run in a paragraph is bold."""
    doc = Document(io.BytesIO(docx_bytes))
    para = doc.paragraphs[para_idx]
    if run_idx >= len(para.runs):
        return False
    return para.runs[run_idx].bold or False


# ── Unit tests: individual edit functions ────────────────────────────────────

class TestApplyReplaceText:
    def test_replaces_text(self):
        doc = Document()
        para = doc.add_paragraph("Original text.")
        _apply_replace_text(para, "New text.")
        assert para.text == "New text."

    def test_preserves_bold_formatting(self):
        """Bold formatting from original run should carry over to replacement."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold text")
        run.bold = True

        _apply_replace_text(para, "Replaced bold text")

        # Check the new run has bold formatting
        assert para.runs[0].bold is True

    def test_preserves_font_size(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Sized text")
        run.font.size = Pt(14)

        _apply_replace_text(para, "New sized text")

        assert para.runs[0].font.size == Pt(14)

    def test_empty_paragraph_handled(self):
        """Paragraph with no runs should not raise."""
        doc = Document()
        para = doc.add_paragraph()
        # Remove all runs manually
        for r in para.runs:
            r._r.getparent().remove(r._r)

        _apply_replace_text(para, "New text")
        # Should not raise — text set via paragraph text property
        assert "New text" in para.text

    def test_multirun_paragraph_becomes_single_run(self):
        """Multi-run paragraph should reduce to single run after replace."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Part 1 ")
        para.add_run("Part 2")
        assert len(para.runs) == 2

        _apply_replace_text(para, "Single replacement")
        assert len(para.runs) == 1
        assert para.runs[0].text == "Single replacement"


class TestApplyReplaceRun:
    def test_replaces_specific_run(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Run 0")
        para.add_run("Run 1")
        para.add_run("Run 2")

        _apply_replace_run(para, run_index=1, new_text="REPLACED")

        assert para.runs[0].text == "Run 0"
        assert para.runs[1].text == "REPLACED"
        assert para.runs[2].text == "Run 2"

    def test_out_of_bounds_run_index_is_silent(self):
        """Out-of-bounds run_index should not raise — just log and skip."""
        doc = Document()
        para = doc.add_paragraph("Single run")
        # run_index=5 is way out of bounds
        _apply_replace_run(para, run_index=5, new_text="should not appear")
        assert para.text == "Single run"  # Unchanged


class TestInsertAfter:
    def test_inserts_paragraph_after_given_index(self):
        doc = Document()
        doc.add_paragraph("Para 0")
        doc.add_paragraph("Para 1")
        doc.add_paragraph("Para 2")

        _insert_paragraph_after(doc, after_index=0, new_text="Inserted after 0")

        texts = [p.text for p in doc.paragraphs]
        assert texts[0] == "Para 0"
        assert texts[1] == "Inserted after 0"
        assert texts[2] == "Para 1"

    def test_insert_after_last_paragraph(self):
        doc = Document()
        doc.add_paragraph("Para 0")
        doc.add_paragraph("Para 1")

        _insert_paragraph_after(doc, after_index=1, new_text="Appended")

        texts = [p.text for p in doc.paragraphs]
        assert "Appended" in texts

    def test_insert_beyond_end_appends(self):
        """Index beyond document length should append without error."""
        doc = Document()
        doc.add_paragraph("Only para")

        _insert_paragraph_after(doc, after_index=999, new_text="Appended safely")

        texts = [p.text for p in doc.paragraphs]
        assert "Appended safely" in texts


class TestDeleteParagraph:
    def test_deletes_middle_paragraph(self):
        doc = Document()
        doc.add_paragraph("Para 0")
        doc.add_paragraph("Para 1 — delete me")
        doc.add_paragraph("Para 2")

        paras = doc.paragraphs
        _delete_paragraph(paras[1])

        remaining = [p.text for p in doc.paragraphs]
        assert "Para 1 — delete me" not in remaining
        assert "Para 0" in remaining
        assert "Para 2" in remaining

    def test_last_paragraph_cleared_not_removed(self):
        """DOCX requires at least one paragraph — last one should be cleared."""
        doc = Document()
        # Add only one paragraph
        para = doc.add_paragraph("Only paragraph")

        _delete_paragraph(para)

        # Document should still have a paragraph (might be empty)
        assert len(doc.paragraphs) >= 1


class TestApplyHeadlineSummary:
    def test_replaces_first_paragraph_as_headline(self):
        doc = Document()
        doc.add_paragraph("Matteo Mantovani")
        doc.add_paragraph("CEO | Founder")
        doc.add_paragraph("Experience section")

        _apply_headline_summary(
            doc=doc,
            new_headline="Senior Strategy Executive",
            new_summary=None,
        )

        assert doc.paragraphs[0].text == "Senior Strategy Executive"
        assert doc.paragraphs[1].text == "CEO | Founder"  # Unchanged

    def test_replaces_second_paragraph_as_summary(self):
        doc = Document()
        doc.add_paragraph("Matteo Mantovani")
        doc.add_paragraph("Old summary text")
        doc.add_paragraph("Experience section")

        _apply_headline_summary(
            doc=doc,
            new_headline=None,
            new_summary="Value creation executive with 10 years...",
        )

        assert doc.paragraphs[0].text == "Matteo Mantovani"  # Unchanged
        assert doc.paragraphs[1].text == "Value creation executive with 10 years..."

    def test_skips_empty_paragraphs_when_finding_headline(self):
        """Empty paragraphs at the top should be skipped."""
        doc = Document()
        doc.add_paragraph("")           # Empty — skip
        doc.add_paragraph("")           # Empty — skip
        doc.add_paragraph("Matteo Mantovani")  # ← headline
        doc.add_paragraph("Old summary")       # ← summary

        _apply_headline_summary(
            doc=doc,
            new_headline="New Headline",
            new_summary="New Summary",
        )

        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "New Headline" in texts
        assert "New Summary" in texts

    def test_both_none_is_noop(self):
        doc = Document()
        doc.add_paragraph("Original")
        _apply_headline_summary(doc=doc, new_headline=None, new_summary=None)
        assert doc.paragraphs[0].text == "Original"


# ── Integration tests: DOCXRenderer ─────────────────────────────────────────

class TestDOCXRenderer:
    """Full round-trip tests applying EditPlan dicts via DOCXRenderer."""

    def test_replace_text_operation(self):
        original = make_docx_bytes(
            "Matteo Mantovani",
            "CEO | Tech Founder",
            "EXPERIENCE",
            "Built Iraq's first super-app.",
        )
        edit_plan = {
            "headline_summary": None,
            "paragraph_edits": [
                {
                    "paragraph_index": 1,
                    "operation": "replace_text",
                    "new_text": "Senior Strategy & Value Creation Executive",
                    "rationale": "Align with e& JD",
                }
            ],
            "keyword_additions": [],
            "sections_to_add": [],
        }

        renderer = DOCXRenderer()
        result = renderer.render(original, edit_plan)

        texts = docx_paragraphs(result.docx_bytes)
        assert "Senior Strategy & Value Creation Executive" in texts
        assert result.edits_applied >= 1

    def test_headline_summary_edit(self):
        original = make_docx_bytes(
            "Matteo Mantovani",
            "CEO | Founder | INSEAD",
            "EXPERIENCE",
        )
        edit_plan = {
            "headline_summary": {
                "new_headline": "Entrepreneur | Founder | Product Builder",
                "new_summary": "I've spent the last decade building things from scratch.",
                "rationale": "Revolut EiR positioning",
            },
            "paragraph_edits": [],
            "keyword_additions": [],
            "sections_to_add": [],
        }

        renderer = DOCXRenderer()
        result = renderer.render(original, edit_plan)

        texts = docx_paragraphs(result.docx_bytes)
        assert "Entrepreneur | Founder | Product Builder" in texts
        assert any("I've spent the last decade" in t for t in texts)

    def test_insert_after_operation(self):
        original = make_docx_bytes(
            "Name",
            "Headline",
            "EXPERIENCE",
            "Old job",
        )
        edit_plan = {
            "headline_summary": None,
            "paragraph_edits": [
                {
                    "paragraph_index": 3,
                    "operation": "insert_after",
                    "new_text": "• Drove 8-15X valuation increase across 4 verticals.",
                    "rationale": "Add missing achievement",
                }
            ],
            "keyword_additions": [],
            "sections_to_add": [],
        }

        renderer = DOCXRenderer()
        result = renderer.render(original, edit_plan)

        texts = docx_paragraphs(result.docx_bytes)
        assert any("8-15X valuation" in t for t in texts)

    def test_delete_operation(self):
        original = make_docx_bytes(
            "Name",
            "Headline",
            "Irrelevant hobby: competitive origami",
            "EXPERIENCE",
        )
        edit_plan = {
            "headline_summary": None,
            "paragraph_edits": [
                {
                    "paragraph_index": 2,
                    "operation": "delete",
                    "rationale": "Not relevant to this role",
                }
            ],
            "keyword_additions": [],
            "sections_to_add": [],
        }

        renderer = DOCXRenderer()
        result = renderer.render(original, edit_plan)

        texts = docx_paragraphs(result.docx_bytes)
        assert "competitive origami" not in texts

    def test_invalid_docx_raises_value_error(self):
        with pytest.raises(ValueError, match="Could not open DOCX"):
            DOCXRenderer().render(b"not a docx", edit_plan={})

    def test_out_of_bounds_paragraph_index_is_skipped(self):
        """Edit targeting non-existent index should be skipped, not crash."""
        original = make_docx_bytes("Only one paragraph")
        edit_plan = {
            "headline_summary": None,
            "paragraph_edits": [
                {
                    "paragraph_index": 999,
                    "operation": "replace_text",
                    "new_text": "Won't apply",
                }
            ],
            "keyword_additions": [],
            "sections_to_add": [],
        }

        renderer = DOCXRenderer()
        result = renderer.render(original, edit_plan)

        # Should succeed with the edit skipped
        assert result.edits_skipped >= 1
        assert result.edits_applied == 0
        # Original content preserved
        assert "Only one paragraph" in docx_paragraphs(result.docx_bytes)

    def test_empty_edit_plan_returns_unchanged_document(self):
        original = make_docx_bytes("Para 1", "Para 2", "Para 3")
        edit_plan = {
            "headline_summary": None,
            "paragraph_edits": [],
            "keyword_additions": [],
            "sections_to_add": [],
        }

        renderer = DOCXRenderer()
        result = renderer.render(original, edit_plan)

        texts = docx_paragraphs(result.docx_bytes)
        assert "Para 1" in texts
        assert "Para 2" in texts
        assert "Para 3" in texts

    def test_sections_to_add(self):
        original = make_docx_bytes("Name", "Summary")
        edit_plan = {
            "headline_summary": None,
            "paragraph_edits": [],
            "keyword_additions": [],
            "sections_to_add": [
                {
                    "heading": "KEY ACHIEVEMENTS",
                    "content": "8-15X valuation increase at Baly.",
                }
            ],
        }

        renderer = DOCXRenderer()
        result = renderer.render(original, edit_plan)

        texts = docx_paragraphs(result.docx_bytes)
        assert "KEY ACHIEVEMENTS" in texts
        assert any("8-15X valuation" in t for t in texts)

    def test_formatting_preserved_across_edits(self):
        """Bold run in original should remain bold after replace_text."""
        original = make_docx_bytes(("Bold content", "Normal", True))
        edit_plan = {
            "headline_summary": None,
            "paragraph_edits": [
                {
                    "paragraph_index": 0,
                    "operation": "replace_text",
                    "new_text": "New bold content",
                }
            ],
            "keyword_additions": [],
            "sections_to_add": [],
        }

        renderer = DOCXRenderer()
        result = renderer.render(original, edit_plan)

        assert docx_run_is_bold(result.docx_bytes, para_idx=0) is True

    def test_multiple_edits_all_applied(self):
        original = make_docx_bytes("Name", "Headline", "Summary", "EXPERIENCE", "Job 1")
        edit_plan = {
            "headline_summary": {
                "new_headline": "New Name",
                "new_summary": None,
                "rationale": "",
            },
            "paragraph_edits": [
                {"paragraph_index": 2, "operation": "replace_text", "new_text": "New summary"},
                {"paragraph_index": 4, "operation": "replace_text", "new_text": "Job 1 rewritten"},
            ],
            "keyword_additions": [],
            "sections_to_add": [],
        }

        renderer = DOCXRenderer()
        result = renderer.render(original, edit_plan)

        texts = docx_paragraphs(result.docx_bytes)
        assert "New Name" in texts
        assert "New summary" in texts
        assert "Job 1 rewritten" in texts
        assert result.edits_applied >= 3

    def test_render_result_metadata(self):
        original = make_docx_bytes("Para 1")
        result = DOCXRenderer().render(original, {
            "headline_summary": None,
            "paragraph_edits": [],
            "keyword_additions": [],
            "sections_to_add": [],
        })

        meta = result.to_metadata()
        assert "edits_applied" in meta
        assert "edits_skipped" in meta
        assert "output_size_bytes" in meta
        assert meta["output_size_bytes"] > 0


# ── render_docx task tests ────────────────────────────────────────────────────

class TestRenderDocxTask:
    """Tests for the render_docx_task with mocked S3 and DB."""

    def _make_mock_job_run(self):
        from app.models.job_run import JobRunStatus
        job_run = MagicMock()
        job_run.id = uuid.uuid4()
        job_run.cv_id = uuid.uuid4()
        job_run.user_id = "test_user"
        job_run.status = JobRunStatus.RENDERING
        job_run.edit_plan = {
            "headline_summary": {
                "new_headline": "Senior Strategy Executive",
                "new_summary": None,
                "rationale": "",
            },
            "paragraph_edits": [],
            "keyword_additions": [],
            "sections_to_add": [],
        }
        job_run.output_s3_key = None
        job_run.is_terminal = False
        job_run.role_title = "Strategy Director"
        job_run.company_name = "TestCo"
        job_run.keyword_match_score = None
        return job_run

    def _make_mock_cv(self):
        cv = MagicMock()
        cv.id = uuid.uuid4()
        cv.s3_key = "cvs/user/abc/cv.docx"
        cv.original_filename = "matteo_cv.docx"
        cv.build_mode = None
        cv.template_slug = None
        return cv

    def test_invalid_uuid_raises(self):
        from app.workers.tasks.render_docx import render_docx_task
        with pytest.raises(ValueError):
            render_docx_task("not-a-uuid")

    @patch("app.workers.tasks.render_docx._complete_render_step")
    @patch("app.services.billing.billing_service.BillingService", new_callable=MagicMock)
    @patch("app.services.email.notifications.notify_pack_complete")
    @patch("app.services.timeline.log_event_sync")
    @patch("app.services.ingest.storage.build_output_s3_key", return_value="outputs/user/run/tailored_cv.docx")
    @patch("app.workers.tasks.render_docx._create_render_step")
    @patch("app.services.ingest.storage.S3StorageService.upload_bytes")
    @patch("app.services.ingest.storage.S3StorageService.download_bytes")
    @patch("app.workers.tasks.render_docx.get_sync_db")
    def test_successful_render_marks_completed(
        self,
        mock_get_db,
        mock_download,
        mock_upload,
        mock_create_step,
        mock_build_key,
        mock_log_event,
        mock_notify,
        mock_billing,
        mock_complete_step,
    ):
        """Happy path: download → render → upload → COMPLETED."""
        from app.models.job_run import JobRunStatus

        job_run = self._make_mock_job_run()
        cv = self._make_mock_cv()

        mock_db = MagicMock()
        mock_db.get.side_effect = lambda model, id: (
            job_run if "JobRun" in model.__name__ else cv
        )
        mock_get_db.return_value.__enter__ = MagicMock(return_value=mock_db)
        mock_get_db.return_value.__exit__ = MagicMock(return_value=False)

        mock_create_step.return_value = uuid.uuid4()

        # Provide a real minimal DOCX for the renderer to process
        mock_download.return_value = make_docx_bytes(
            "Matteo Mantovani",
            "CEO | Founder",
            "EXPERIENCE",
        )
        mock_upload.return_value = "outputs/user/run/tailored_cv.docx"

        from app.workers.tasks.render_docx import render_docx_task
        result = render_docx_task(str(job_run.id))

        assert result["status"] == "completed"
        assert "output_s3_key" in result
        mock_upload.assert_called_once()

    @patch("app.workers.tasks.render_docx._mark_run_failed")
    @patch("app.workers.tasks.render_docx._fail_render_step")
    @patch("app.workers.tasks.render_docx._create_render_step")
    @patch("app.services.ingest.storage.S3StorageService.download_bytes")
    @patch("app.workers.tasks.render_docx.get_sync_db")
    def test_s3_download_failure_triggers_retry(
        self,
        mock_get_db,
        mock_download,
        mock_create_step,
        mock_fail_step,
        mock_mark_failed,
    ):
        """S3 download failure should trigger retry and eventually mark failed."""
        from app.api.middleware.error_handler import StorageError

        job_run = self._make_mock_job_run()
        cv = self._make_mock_cv()

        mock_db = MagicMock()
        mock_db.get.side_effect = lambda model, id: (
            job_run if "JobRun" in model.__name__ else cv
        )
        mock_get_db.return_value.__enter__ = MagicMock(return_value=mock_db)
        mock_get_db.return_value.__exit__ = MagicMock(return_value=False)

        mock_create_step.return_value = uuid.uuid4()
        mock_download.side_effect = StorageError("S3 unavailable")

        from app.workers.tasks.render_docx import render_docx_task

        # Task should raise on max retries exceeded
        with pytest.raises(Exception):
            render_docx_task(str(job_run.id))

        mock_fail_step.assert_called()
