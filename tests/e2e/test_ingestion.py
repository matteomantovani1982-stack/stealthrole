"""
tests/e2e/test_ingestion.py

End-to-end tests for POST /api/v1/cvs/ingest.
"""

import uuid
import io

import httpx
import pytest

BASE_URL = "http://localhost:8000"


@pytest.fixture(scope="module")
def auth_headers():
    """Register + login, return auth headers."""
    email = f"ingest_test_{uuid.uuid4().hex[:8]}@test.com"
    password = "TestPass123!"

    httpx.post(f"{BASE_URL}/api/v1/auth/register", json={
        "email": email,
        "password": password,
        "full_name": "Ingest Tester",
    })

    r = httpx.post(f"{BASE_URL}/api/v1/auth/login", json={
        "email": email,
        "password": password,
    })
    assert r.status_code == 200
    token = r.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


class TestIngestUnauth:
    """POST /api/v1/cvs/ingest without auth."""

    def test_ingest_returns_401_without_auth(self):
        r = httpx.post(f"{BASE_URL}/api/v1/cvs/ingest")
        assert r.status_code in (401, 403)


class TestIngestValidation:
    """POST /api/v1/cvs/ingest with invalid inputs."""

    def test_ingest_returns_422_no_file(self, auth_headers):
        """Should return 422 when no file is provided."""
        r = httpx.post(f"{BASE_URL}/api/v1/cvs/ingest", headers=auth_headers)
        assert r.status_code == 422

    def test_ingest_returns_422_unsupported_type(self, auth_headers):
        """Should return 422 for unsupported file types."""
        files = {"file": ("test.txt", b"some text content", "text/plain")}
        r = httpx.post(
            f"{BASE_URL}/api/v1/cvs/ingest",
            headers=auth_headers,
            files=files,
        )
        assert r.status_code == 422


class TestIngestSuccess:
    """POST /api/v1/cvs/ingest with valid DOCX-like input."""

    def test_ingest_docx_returns_correct_shape(self, auth_headers):
        """Should return document_type, anomaly_flags, and file_id."""
        # Create a minimal DOCX-like file — the endpoint will try to parse it
        # and may fail with 422 if python-docx can't read it, which is acceptable.
        # We test the shape if it succeeds, or accept 422.
        from docx import Document
        doc = Document()
        doc.add_paragraph("Experience")
        doc.add_paragraph("Software Engineer at Acme Corp")
        doc.add_paragraph("Education")
        doc.add_paragraph("BS Computer Science")
        doc.add_paragraph("Skills: Python, FastAPI")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        files = {"file": ("test_cv.docx", buf.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = httpx.post(
            f"{BASE_URL}/api/v1/cvs/ingest",
            headers=auth_headers,
            files=files,
        )
        assert r.status_code == 200
        data = r.json()
        assert "document_type" in data
        assert data["document_type"] in ("cv", "jd", "other")
        assert "anomaly_flags" in data
        assert isinstance(data["anomaly_flags"], list)
        assert "file_id" in data
        assert len(data["file_id"]) > 0
