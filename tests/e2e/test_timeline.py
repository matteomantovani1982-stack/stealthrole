"""
tests/e2e/test_timeline.py

End-to-end tests for GET /api/v1/jobs/{job_run_id}/timeline.
"""

import uuid

import httpx
import pytest

BASE_URL = "http://localhost:8000"


@pytest.fixture(scope="module")
def auth_headers():
    """Register + login, return auth headers."""
    email = f"timeline_test_{uuid.uuid4().hex[:8]}@test.com"
    password = "TestPass123!"

    httpx.post(f"{BASE_URL}/api/v1/auth/register", json={
        "email": email,
        "password": password,
        "full_name": "Timeline Tester",
    })

    r = httpx.post(f"{BASE_URL}/api/v1/auth/login", json={
        "email": email,
        "password": password,
    })
    assert r.status_code == 200
    token = r.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


class TestTimelineUnauth:
    """GET /api/v1/jobs/{id}/timeline without auth."""

    def test_timeline_returns_401_without_auth(self):
        fake_id = str(uuid.uuid4())
        r = httpx.get(f"{BASE_URL}/api/v1/jobs/{fake_id}/timeline")
        assert r.status_code in (401, 403)


class TestTimeline404:
    """GET /api/v1/jobs/{id}/timeline with non-existent job run."""

    def test_timeline_returns_404_for_fake_id(self, auth_headers):
        fake_id = str(uuid.uuid4())
        r = httpx.get(
            f"{BASE_URL}/api/v1/jobs/{fake_id}/timeline",
            headers=auth_headers,
        )
        assert r.status_code == 404


class TestTimelineShape:
    """GET /api/v1/jobs/{id}/timeline response shape."""

    def test_timeline_returns_200_with_correct_shape(self, auth_headers):
        """Create a job run first, then verify timeline shape."""
        # First, we need a CV. Upload one.
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("Experience: Software Engineer at Acme Corp, 2020-2024")
        doc.add_paragraph("Education: BS Computer Science")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        upload_r = httpx.post(
            f"{BASE_URL}/api/v1/cvs",
            headers=auth_headers,
            files={"file": ("timeline_test.docx", buf.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        if upload_r.status_code not in (200, 201, 202):
            pytest.skip(f"CV upload failed with {upload_r.status_code}, skipping timeline shape test")

        cv_id = upload_r.json().get("cv_id") or upload_r.json().get("id")
        if not cv_id:
            pytest.skip("Could not get cv_id from upload response")

        # Create a job run
        job_r = httpx.post(
            f"{BASE_URL}/api/v1/jobs",
            headers=auth_headers,
            json={
                "cv_id": cv_id,
                "jd_text": "We are looking for a software engineer with Python experience.",
            },
        )
        if job_r.status_code not in (200, 201, 202):
            pytest.skip(f"Job run creation failed with {job_r.status_code}")

        job_run_id = job_r.json().get("job_run_id") or job_r.json().get("id")
        if not job_run_id:
            pytest.skip("Could not get job_run_id from job response")

        # Get timeline
        r = httpx.get(
            f"{BASE_URL}/api/v1/jobs/{job_run_id}/timeline",
            headers=auth_headers,
        )
        assert r.status_code == 200
        data = r.json()
        assert "events" in data
        assert "total" in data
        assert isinstance(data["events"], list)
        assert isinstance(data["total"], int)
